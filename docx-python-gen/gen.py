import json
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx2pdf import convert

def generate_word_document(data):
    # Create a new Word document
    document = Document()

    # Add a heading
    document.add_heading('JSON Data', level=1)

    # Process each item in the JSON data
    for item in data:
        # Add a heading for each item
        document.add_heading(item['title'], level=2)

        # Add a table for the item's details
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Set column widths
        table.autofit = False
        table.columns[0].width = Pt(200)
        table.columns[1].width = Pt(400)

        # Add details to the table
        for key, value in item['details'].items():
            row = table.add_row().cells
            row[0].text = key
            row[1].text = value

        # Add spacing between items
        document.add_paragraph()

    # Save the document
    document.save('output.docx')
    convert('output.docx')

# Example JSON data
json_data = [
    {
        'title': 'Item 1',
        'details': {
            'Property 1': 'Value 1',
            'Property 2': 'Value 2',
            'Property 3': 'Value 3'
        }
    },
    {
        'title': 'Item 2',
        'details': {
            'Property 1': 'Value 1',
            'Property 2': 'Value 2',
            'Property 3': 'Value 3'
        }
    }
]

# Generate the Word document
generate_word_document(json_data)

